"""
File Parser Module
Handles text extraction from PDF and DOCX files with advanced preprocessing
"""
import os
import re
import PyPDF2
from docx import Document
from typing import Optional, List
import tempfile

class ResumeParser:
    """Advanced resume parser with text extraction and preprocessing"""
    
    def __init__(self):
        self.sections = [
            'education', 'experience', 'skills', 'projects',
            'certifications', 'languages', 'publications', 'awards'
        ]
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extract text from PDF or DOCX file based on file extension
        """
        if file_extension == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == 'docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file with error handling
        """
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    try:
                        pdf_reader.decrypt('')
                    except:
                        raise Exception("Encrypted PDF cannot be processed")
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                if not text.strip():
                    raise Exception("No text could be extracted from PDF")
                    
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
        
        return self._clean_text(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file with paragraph handling
        """
        text = ""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                raise Exception("No text could be extracted from DOCX")
                
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        """
        # Replace multiple newlines with single newline
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\-\+\@\#]', '', text)
        
        # Remove email addresses (optional - can be kept for analysis)
        # text = re.sub(r'\S+@\S+', '[EMAIL]', text)
        
        # Remove URLs
        text = re.sub(r'http\S+|www\S+', '[URL]', text)
        
        # Remove phone numbers (optional)
        # text = re.sub(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', '[PHONE]', text)
        
        return text.strip()
    
    def extract_sections(self, text: str) -> dict:
        """
        Attempt to extract different sections from resume text
        This is a simple implementation - can be enhanced with ML/NLP
        """
        sections = {}
        lines = text.split('\n')
        current_section = 'other'
        section_content = []
        
        # Common section headers in resumes
        section_headers = {
            'education': ['education', 'academic background', 'qualifications'],
            'experience': ['experience', 'work experience', 'employment', 'work history'],
            'skills': ['skills', 'technical skills', 'core competencies'],
            'projects': ['projects', 'personal projects', 'academic projects'],
            'certifications': ['certifications', 'certificates', 'professional development'],
            'languages': ['languages', 'language proficiency'],
            'publications': ['publications', 'papers', 'research'],
            'awards': ['awards', 'honors', 'achievements']
        }
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            for section, keywords in section_headers.items():
                if any(keyword in line_lower for keyword in keywords) and len(line) < 50:
                    # Save previous section
                    if current_section != 'other' and section_content:
                        sections[current_section] = '\n'.join(section_content)
                    
                    # Start new section
                    current_section = section
                    section_content = []
                    break
            else:
                # Not a section header, add to current section
                if line.strip():
                    section_content.append(line)
        
        # Save last section
        if current_section != 'other' and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        # Add full text as well
        sections['full_text'] = text
        
        return sections


class BatchResumeParser:
    """Handle multiple resume parsing operations"""
    
    def __init__(self):
        self.parser = ResumeParser()
    
    def parse_batch(self, file_paths: List[str]) -> List[dict]:
        """
        Parse multiple resumes and return their text and metadata
        """
        results = []
        for file_path in file_paths:
            try:
                file_extension = file_path.split('.')[-1].lower()
                text = self.parser.extract_text(file_path, file_extension)
                sections = self.parser.extract_sections(text)
                
                results.append({
                    'filename': os.path.basename(file_path),
                    'text': text,
                    'sections': sections,
                    'success': True
                })
            except Exception as e:
                results.append({
                    'filename': os.path.basename(file_path),
                    'error': str(e),
                    'success': False
                })
        
        return results


# Utility function for quick text extraction
def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Quick utility function to extract text from a file
    """
    parser = ResumeParser()
    try:
        file_extension = file_path.split('.')[-1].lower()
        return parser.extract_text(file_path, file_extension)
    except Exception as e:
        print(f"Error extracting text: {e}")
        return None


# Temporary file handler
def save_uploaded_file(uploaded_file) -> str:
    """
    Save an uploaded file to temporary location and return the path
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=uploaded_file.filename) as tmp_file:
        uploaded_file.save(tmp_file.name)
        return tmp_file.name


def cleanup_temp_file(file_path: str):
    """
    Remove temporary file
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        print(f"Error cleaning up temp file: {e}")